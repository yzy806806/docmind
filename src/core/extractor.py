"""Text extraction from various file formats.

Provides synchronous extraction for .txt, .md, .pdf, .docx, .html/.htm.
Includes page-by-page PDF extraction, size-tiered ProcessPoolExecutor routing,
memory estimation, markdown-to-plaintext conversion, and hash utilities.
"""

import concurrent.futures
import hashlib
import os
import re
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# Size-tiered ProcessPoolExecutor routing
# ---------------------------------------------------------------------------

NORMAL_MAX_WORKERS = 4
LARGE_MAX_WORKERS = 1
LARGE_FILE_THRESHOLD = 50 * 1024 * 1024  # 50 MB

_normal_pool: concurrent.futures.ProcessPoolExecutor | None = None
_large_pool: concurrent.futures.ProcessPoolExecutor | None = None


def _get_pool(
    file_size: int,
) -> concurrent.futures.ProcessPoolExecutor:
    """Return the appropriate pool based on file size (lazy init, module-level singletons)."""
    global _normal_pool, _large_pool

    if file_size >= LARGE_FILE_THRESHOLD:
        if _large_pool is None:
            _large_pool = concurrent.futures.ProcessPoolExecutor(
                max_workers=LARGE_MAX_WORKERS
            )
        return _large_pool
    else:
        if _normal_pool is None:
            _normal_pool = concurrent.futures.ProcessPoolExecutor(
                max_workers=NORMAL_MAX_WORKERS
            )
        return _normal_pool


def shutdown_pools(wait: bool = True) -> None:
    """Shut down both executor pools.  Safe to call multiple times."""
    global _normal_pool, _large_pool
    if _normal_pool is not None:
        _normal_pool.shutdown(wait=wait)
        _normal_pool = None
    if _large_pool is not None:
        _large_pool.shutdown(wait=wait)
        _large_pool = None


# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------

def file_sha256(file_path: Path) -> str:
    """Compute SHA-256 hash of a file on disk."""
    sha = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            sha.update(chunk)
    return sha.hexdigest()


# Memory multiplier by file type (used for pre-scan estimation)
_MEMORY_MULTIPLIER: dict[str, int] = {
    ".pdf": 10,
    ".docx": 5,
    ".html": 3,
    ".htm": 3,
}


def estimate_memory(file_path: Path) -> int:
    """Estimate peak memory (bytes) needed to extract *file_path*.

    The estimate is ``file_size * multiplier``, where the multiplier
    depends on the file extension.  Falls back to 2× for unknown types.
    """
    try:
        size = file_path.stat().st_size
    except OSError:
        return 0
    ext = file_path.suffix.lower()
    multiplier = _MEMORY_MULTIPLIER.get(ext, 2)
    return size * multiplier


def check_memory_budget(
    file_path: Path, max_memory: int = 512 * 1024 * 1024
) -> bool:
    """Return ``True`` if the estimated memory for *file_path* is within *max_memory*."""
    return estimate_memory(file_path) <= max_memory


# ---------------------------------------------------------------------------
# Markdown-to-plaintext (regex-based, no external markdown lib)
# ---------------------------------------------------------------------------

# Pattern to match fenced code blocks (``` ... ```).
_FENCED_RE = re.compile(r"```[^`]*```", re.DOTALL)

# Inline code: single or double backtick spans.
_INLINE_CODE_RE = re.compile(r"`{1,2}[^`\n]+?`{1,2}")

# ATX headings: optional leading spaces, 1-6 #, optional space, then text.
_HEADING_RE = re.compile(r"^\s*#{1,6}\s+(.+?)\s*#*\s*$", re.MULTILINE)

# Link syntax: [text](url) — keep text.
_LINK_RE = re.compile(r"\[([^\]]*)\]\([^)]*\)")

# Image syntax: ![alt](url) — remove entirely.
_IMAGE_RE = re.compile(r"!\[[^\]]*\]\([^)]*\)")


def _md_to_plaintext(text: str) -> str:
    """Convert Markdown text to plain text using regex substitutions."""
    # Remove fenced code blocks first (they can contain anything).
    text = _FENCED_RE.sub("", text)
    # Remove inline code spans.
    text = _INLINE_CODE_RE.sub("", text)
    # Convert ATX headings to just the heading text.
    text = _HEADING_RE.sub(r"\1", text)
    # Remove image syntax.
    text = _IMAGE_RE.sub("", text)
    # Replace link syntax with just the link text.
    text = _LINK_RE.sub(r"\1", text)
    # Collapse runs of blank lines.
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Extractor class
# ---------------------------------------------------------------------------

class Extractor:
    """Extract plain text from supported file types.

    Supported extensions: .txt, .md, .pdf, .docx, .html, .htm
    """

    SUPPORTED: set[str] = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    @staticmethod
    def extract(file_path: Path) -> Optional[str]:
        """Extract text from a file, dispatching by extension.

        Returns the full extracted text as a string, or ``None`` on failure
        or when the extension is unsupported.
        """
        ext = file_path.suffix.lower()
        if ext not in Extractor.SUPPORTED:
            return None

        try:
            if ext == ".pdf":
                return Extractor._extract_pdf(file_path)
            elif ext == ".docx":
                return Extractor._extract_docx(file_path)
            elif ext in (".html", ".htm"):
                return Extractor._extract_html(file_path)
            elif ext == ".md":
                return Extractor._extract_markdown(file_path)
            else:  # .txt
                return file_path.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            print(f"[Extractor] Failed to extract {file_path}: {e}")
            return None

    # ------------------------------------------------------------------
    # Static helpers delegated to module-level functions
    # ------------------------------------------------------------------

    @staticmethod
    def estimate_memory(file_path: Path) -> int:
        """Estimate peak memory needed to extract *file_path*."""
        return estimate_memory(file_path)

    @staticmethod
    def check_memory_budget(file_path: Path, max_memory: int = 512 * 1024 * 1024) -> bool:
        """Return True if estimated memory is within budget."""
        return check_memory_budget(file_path, max_memory)

    @staticmethod
    def file_sha256(file_path: Path) -> str:
        """Compute SHA-256 hash of a file on disk."""
        return file_sha256(file_path)

    @staticmethod
    def shutdown_pools(wait: bool = True) -> None:
        """Shut down both ProcessPoolExecutor pools."""
        shutdown_pools(wait=wait)

    # ------------------------------------------------------------------

    @staticmethod
    def extract_from_bytes(content: bytes, ext: str) -> Optional[str]:
        """Extract text from in-memory bytes (e.g. from WebDAV/streaming).

        *ext* must include the leading dot, e.g. ``".pdf"``.
        """
        ext = ext.lower()
        if ext not in Extractor.SUPPORTED:
            return None

        try:
            if ext == ".pdf":
                return Extractor._extract_pdf_bytes(content)
            elif ext == ".docx":
                return Extractor._extract_docx_bytes(content)
            elif ext in (".html", ".htm"):
                return Extractor._extract_html_bytes(content)
            elif ext == ".md":
                text = content.decode("utf-8", errors="replace")
                return _md_to_plaintext(text)
            else:  # .txt
                return content.decode("utf-8", errors="replace")
        except Exception as e:
            print(f"[Extractor] Failed to extract from bytes ({ext}): {e}")
            return None

    @staticmethod
    def extract_pages(file_path: Path) -> list[tuple[int, str]]:
        """Extract PDF text page by page.

        Returns a list of ``(page_number, page_text)`` tuples.
        Page numbers are 1-based.  Returns an empty list on failure.
        """
        ext = file_path.suffix.lower()
        if ext != ".pdf":
            return []

        try:
            import pdfplumber

            pages: list[tuple[int, str]] = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ""
                    pages.append((page_num, page_text))
            return pages
        except Exception as e:
            print(f"[Extractor] Failed to extract pages from {file_path}: {e}")
            return []

    @staticmethod
    def extract_async(
        file_path: Path,
    ) -> concurrent.futures.Future:
        """Submit extraction to the appropriate ProcessPoolExecutor.

        Files >= 50 MB are routed to a single-worker *large* pool;
        smaller files use a 4-worker *normal* pool.

        Returns a :class:`~concurrent.futures.Future` that resolves to
        the extracted text (``str``) or ``None``.
        """
        try:
            file_size = file_path.stat().st_size
        except OSError:
            file_size = 0

        pool = _get_pool(file_size)
        return pool.submit(Extractor.extract, file_path)

    # ------------------------------------------------------------------
    # PDF extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_pdf(file_path: Path) -> str:
        import pdfplumber

        text_parts: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_pdf_bytes(content: bytes) -> str:
        import io

        import pdfplumber

        text_parts: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    # ------------------------------------------------------------------
    # DOCX extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_docx(file_path: Path) -> str:
        from docx import Document

        doc = Document(str(file_path))
        parts: list[str] = []

        # Paragraph text
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)

        # Table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text for cell in row.cells if cell.text
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    @staticmethod
    def _extract_docx_bytes(content: bytes) -> str:
        import io

        from docx import Document

        doc = Document(io.BytesIO(content))
        parts: list[str] = []

        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text for cell in row.cells if cell.text
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    # ------------------------------------------------------------------
    # HTML extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_html(file_path: Path) -> str:
        from bs4 import BeautifulSoup

        raw = file_path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    @staticmethod
    def _extract_html_bytes(content: bytes) -> str:
        from bs4 import BeautifulSoup

        raw = content.decode("utf-8", errors="replace")
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    # ------------------------------------------------------------------
    # Markdown extraction
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_markdown(file_path: Path) -> str:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return _md_to_plaintext(text)
